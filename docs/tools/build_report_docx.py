#!/usr/bin/env python3
"""
프로그램명 : build_report_docx.py — Day 2 종합실습 제출 보고서(docx) 생성
작성자     : 판교 10반 박유진
작성일     : 2026-08-08
설명       : main.py 실행 결과와 캡처 이미지를 모아 제출용 Word 보고서를 만든다.
             종합실습4 보고서 양식(제목·목차·Heading·Code Block 스타일)을 템플릿으로
             재사용하므로 서식이 이전 제출물과 동일하게 유지된다.
변경 이력  : 2026-08-08 최초 작성
실행 방법  : <toolvenv>/bin/python docs/tools/build_report_docx.py
             (python-docx 필요. 프로젝트 requirements.txt 와는 무관한 문서 생성 전용 도구)
산출 파일  : docs/판교_10반_박유진_day2종합실습.docx
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
CDIR = os.path.join(BASE, "docs", "captures")
FIGDIR = os.path.join(BASE, "outputs", "figures")
TEMPLATE = ("/Users/yujin/projects/skala-workspace/0730_skala_assignment/"
            "종합실습4보고서_양식.docx")
OUT = os.path.join(BASE, "docs", "판교_10반_박유진_day2종합실습.docx")

TITLE = "End2End 데이터 분석 프로젝트 — Day 2 종합 실습"
SUBTITLE = "NYC 옐로우캡 운행 기록 409만 건으로 만든 정체 예측 파이프라인"
AUTHOR = "판교 10반 P323 박유진"

HEAD_BG = "F2F2F2"
BORDER = "BFBFBF"
GRAY = RGBColor(0x77, 0x77, 0x77)
IMG_W = Cm(15.5)


# ═══════════════════════════════════════════════════════════════
# 서식 헬퍼 — 종합실습4 빌더와 동일한 서식을 쓰기 위해 같은 구현을 사용한다
# ═══════════════════════════════════════════════════════════════
def shade(el, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    el.append(shd)


def cell_format(cell, fill=None):
    """셀에 실선 테두리와 안쪽 여백을 준다. 템플릿 표 스타일에 테두리가 없어 직접 그린다."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), BORDER)
        borders.append(b)
    tcPr.append(borders)
    if fill:
        shade(tcPr, fill)
    mar = OxmlElement("w:tcMar")
    for side, w in (("top", 60), ("left", 90), ("bottom", 60), ("right", 90)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(w))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def cell_text(cell, text, bold=False, size=9):
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    return p


def make_table(doc, header, rows, widths=None, size=9, bold_first=False):
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=len(rows[0]))
    t.autofit = True
    ri = 0
    if header:
        for j, hd in enumerate(header):
            c = t.cell(0, j)
            cell_format(c, HEAD_BG)
            cell_text(c, hd, bold=True, size=size)
            if widths:
                c.width = Cm(widths[j])
        ri = 1
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.cell(i + ri, j)
            cell_format(c, HEAD_BG if (bold_first and j == 0) else None)
            cell_text(c, str(v), bold=(bold_first and j == 0), size=size)
            if widths:
                c.width = Cm(widths[j])
    doc.add_paragraph()
    return t


def add_code(doc, code):
    p = doc.add_paragraph(style="Code Block")
    lines = code.split("\n")
    for i, line in enumerate(lines):
        r = p.add_run(line if line else " ")
        if i < len(lines) - 1:
            r.add_break()
    return p


def add_toc(doc):
    """워드 목차 필드를 삽입한다. 문서를 열고 F9 를 누르면 실제 목차로 채워진다."""
    p = doc.add_paragraph()
    r = p.add_run()
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText")
    i.set(qn("xml:space"), "preserve")
    i.text = r'TOC \o "1-2" \h \z \u'
    s = OxmlElement("w:fldChar")
    s.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "[ 목차를 클릭한 뒤 F9 를 눌러 갱신하세요 ]"
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), "end")
    for el in (b, i, s, t, e):
        r._r.append(el)
    return p


MISSING = []


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY
    return p


def add_image(doc, path, label, width=IMG_W):
    """이미지를 넣고 아래에 설명을 단다. 파일이 없으면 자리표시자를 남긴다."""
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption(doc, label)
        return True
    MISSING.append(os.path.basename(path))
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    cell_format(c, "FAFAFA")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run(f"{label}\n{os.path.basename(path)}")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY
    doc.add_paragraph()
    return False


def cap(doc, fname, label, width=IMG_W):
    return add_image(doc, os.path.join(CDIR, fname), label, width)


def fig(doc, fname, label, width=IMG_W):
    return add_image(doc, os.path.join(FIGDIR, fname), label, width)


def sub(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(10)
    return p


def h(doc, text, level, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    if page_break:
        p.paragraph_format.page_break_before = True
    return p


def body(doc, text, bullet=False):
    p = doc.add_paragraph()
    if bullet:
        p.paragraph_format.left_indent = Cm(0.5)
        text = "· " + text
    p.add_run(text)
    return p


# ═══════════════════════════════════════════════════════════════
# 본문
# ═══════════════════════════════════════════════════════════════
def chapter_overview(doc):
    h(doc, "1. 개요", 1, page_break=True)
    body(doc,
         "NYC 옐로우캡 운행 기록 409만 건을 대상으로, 승차 시점에 알 수 있는 정보만으로 "
         "그 트립이 정체 구간에 걸릴지 예측하는 End2End 파이프라인을 구현했습니다. "
         "데이터 로딩 → 전처리 → EDA → 시각화 → 통계 검정 → 모델 학습 → 보고서 생성의 "
         "일곱 단계로 나누고, python main.py 한 번으로 전 과정이 약 15초에 끝나도록 "
         "구성했습니다.")

    h(doc, "1-1. 분석 가설과 타깃 정의", 2)
    body(doc,
         "가설은 “승차 시각·요일·출발/도착 존·주행거리·승객 수만으로 그 트립의 정체 여부를 "
         "예측할 수 있다”입니다. 소요시간·요금·팁처럼 하차 후에야 확정되는 값은 타깃 누수를 "
         "일으키므로 피처에서 제외했습니다.")
    body(doc,
         "정체의 정의는 하나로 정하지 않고 두 가지를 만들어 함께 학습했습니다. 정의를 바꾸면 "
         "성능과 해석이 어떻게 맞바뀌는지 확인하기 위해서입니다.")
    make_table(doc, ["타깃", "정의", "임계값", "의미"], [
        ["jam_abs 절대정체", "전체 트립 평균속도 하위 33%", "7.05 mph 미만", "절대적으로 느린 트립"],
        ["jam_rel 상대정체", "동일 경로 평소 중앙속도 대비 하위 33%", "평소의 86.4% 미만", "평소보다 막힌 트립"],
    ], widths=[3.4, 5.6, 3.0, 3.5])

    h(doc, "1-2. 사용 데이터셋", 2)
    make_table(doc, None, [
        ["데이터셋", "NYC TLC Yellow Taxi 2026년 5월"],
        ["출처", "https://d37ci6vzurychx.cloudfront.net/trip-data/"
                 "yellow_tripdata_2026-05.parquet"],
        ["원본 규모", "4,090,836행 × 20열 (약 66MB, 압축 Parquet)"],
        ["정제 후", "2,841,553행 × 29열 (30.5% 제거)"],
    ], widths=[3.0, 12.5], bold_first=True)

    h(doc, "1-3. 실행 환경 및 기술 스택", 2)
    make_table(doc, ["구분", "사용 기술"], [
        ["실행 환경", "Python 3.11.15 / venv / macOS (Apple Silicon)"],
        ["데이터 처리", "pandas 3.0.5, polars 1.43.2, pyarrow 25.0.0"],
        ["시각화", "seaborn 0.13.2 + matplotlib 3.11.1, plotly 6.9.0"],
        ["통계", "scipy 1.17.1 (stats.ttest_ind)"],
        ["머신러닝", "scikit-learn 1.9.0 (Pipeline, ColumnTransformer), joblib 1.5.3"],
        ["코드 품질", "ruff (line-length 100, select E/F/I/UP/B)"],
    ], widths=[3.0, 12.5])

    h(doc, "1-4. 폴더 구조와 실행 방법", 2)
    add_code(doc, """skala-python-part2/
├── main.py                  # 전체 파이프라인 실행 (1~7단계)
├── requirements.txt
├── pyproject.toml           # ruff 정적 검사 설정
├── README.md
├── data/
│   └── yellow_tripdata_2026-05.parquet
├── src/
│   ├── config.py            # 경로·상수·컬럼 정의 (모든 설정의 단일 출처)
│   ├── common.py            # 공용 유틸 (구분선, p-value 해석, 마크다운 표)
│   ├── data_loader.py       # [1] Pandas / Polars 로딩 및 결과 대조
│   ├── preprocess.py        # [2] 결측 진단 · 정제 · 타깃 생성
│   ├── eda.py               # [3] 기술통계 · 상관계수
│   ├── visualize.py         # [4] Seaborn 정적 / Plotly 인터랙티브
│   ├── stats_test.py        # [5] t-test 및 p-value 해석
│   ├── model.py             # [6] sklearn Pipeline 학습 · 평가 · 저장
│   └── report.py            # [7] report.md 자동 생성
├── docs/
│   ├── 컬럼_모델_정리.md      # 조원 공유용 컬럼·모델 정리
│   └── captures/            # 실행 결과 캡처
└── outputs/
    ├── report.md            # 자동 생성 보고서 (378줄)
    ├── figures/             # 정적 차트 3종(PNG) + 인터랙티브 2종(HTML)
    ├── models/              # 학습된 파이프라인 2종(joblib)
    └── tables/              # 기술통계·상관계수·t-test·모델 비교표(CSV)""")
    add_code(doc, """# 실행
python -m venv .venv
source .venv/bin/activate
pip install -r requirements.txt
python main.py""")


STAGES = [
    dict(
        num="2-1", title="1단계. 데이터 준비 — Pandas / Polars 이중 로딩",
        cap="cap_01_load.png",
        label="[ 1단계 실행 — 로딩 속도 비교, 결과 일치 검증, 집계 대조 ]",
        paras=[
            "같은 parquet 파일을 세 가지 방식으로 읽어 결과가 일치하는지 검증했습니다. "
            "Polars 는 교안에서 다룬 Eager(read_parquet)와 Lazy(scan_parquet)를 모두 "
            "사용했습니다.",
        ],
        tables=[
            (["방식", "소요 시간", "비고"], [
                ["Pandas read_parquet", "0.248초", "기준"],
                ["Polars read_parquet (Eager)", "0.141초", "Pandas 대비 1.76배"],
                ["Pandas groupby (집계)", "0.216초", "기준"],
                ["Polars scan_parquet + group_by (Lazy)", "0.106초", "Pandas 대비 2.04배"],
            ], [6.5, 3.0, 6.0]),
        ],
        after=[
            "처음에는 Polars 로 읽은 결과를 쓰지 않고 속도만 재고 있었습니다. 그건 비교가 "
            "아니라 시연이라는 생각이 들어, 같은 집계를 두 방식으로 수행해 결과를 대조하도록 "
            "바꿨습니다. 그러자 그룹 수가 8개와 9개로 달랐습니다.",
            "숫자가 있는 그룹은 평균·중앙값까지 소수점 셋째 자리가 전부 일치했고, 유일한 "
            "차이가 NaN 그룹 879,400건이었습니다. Pandas 의 groupby 는 그룹 키가 결측인 행을 "
            "기본적으로 제외하지만, Polars 의 group_by 는 null 을 하나의 그룹으로 유지하기 "
            "때문입니다. 이 null 그룹이 2단계에서 확인할 구조적 결측 블록과 같은 행들입니다.",
        ],
    ),
    dict(
        num="2-2", title="2단계. 전처리 — 결측·중복·이상치 처리 및 타깃 생성",
        cap="cap_02_preprocess.png",
        label="[ 2단계 실행 — 구조적 결측 블록 탐지, 정제 이력, 타깃 생성 ]",
        paras=[
            "결측 수를 세는 것에서 그치지 않고, 결측이 무작위인지 구조적인지 먼저 "
            "판별했습니다. 결측이 있는 5개 컬럼(passenger_count, RatecodeID, "
            "store_and_fwd_flag, congestion_surcharge, Airport_fee)의 결측 위치가 완전히 "
            "일치했고, 그 크기는 955,371행(23.4%)이었습니다.",
        ],
        tables=[
            (["VendorID", "결측 블록", "나머지"], [
                ["1", "123,169", "682,052"],
                ["2", "824,559", "2,401,663"],
                ["6", "7,643", "0"],
                ["7", "0", "51,750"],
            ], [3.0, 4.0, 4.0]),
        ],
        after=[
            "블록 안팎의 벤더 구성이 겹치지 않았습니다. 벤더 6은 결측 블록에만, 벤더 7은 "
            "나머지에만 존재해, 서로 다른 데이터 파이프라인이 한 파일에 섞여 있었음을 "
            "확인했습니다. 이걸 모르고 fillna(mean) 으로 채웠다면 존재하지 않는 값을 지어내며 "
            "두 소스를 하나로 뭉개는 결과가 됐을 것입니다.",
            "또한 isna() 로 잡히지 않는 코드값 결측도 별도로 진단했습니다. "
            "RatecodeID = 99(unknown)는 140,897행(3.44%)으로 ‘미상’ 자체가 정보이고 보조 "
            "피처라 별도 범주로 유지했고, 존 ID 264·265(Unknown/N-A)는 25,260행(0.62%)으로 "
            "존이 핵심 피처라 제거했습니다.",
            "정제는 의존 관계 순서를 따랐습니다. 소요시간을 먼저 걸러야 0으로 나누지 않고 "
            "속도를 계산할 수 있으므로 속도 기반 필터를 뒤에 두었습니다. 최종적으로 "
            "4,090,836행에서 2,997,960행(26.72% 제거)이 되었고, 상대정체 계산에 필요한 "
            "경로 표본 조건까지 적용해 2,841,553행으로 분석을 진행했습니다.",
        ],
    ),
    dict(
        num="2-3", title="3단계. EDA — 기술통계 및 상관계수",
        cap="cap_03_eda.png",
        label="[ 3단계 실행 — 기술통계(평균·표준편차·분위수)와 상관계수 행렬 ]",
        paras=[
            "수치형 7개 변수에 대해 describe() 기반 기술통계와 피어슨 상관계수를 산출했습니다.",
        ],
        tables=[
            (["변수", "평균", "표준편차", "최솟값", "중앙값", "최댓값"], [
                ["trip_distance", "3.057", "3.983", "0.02", "1.640", "47.660"],
                ["duration_min", "17.026", "13.803", "1.00", "13.117", "179.983"],
                ["speed_mph", "9.808", "5.767", "1.00", "8.451", "69.751"],
                ["fare_amount", "18.815", "15.875", "0.01", "13.500", "615.000"],
                ["tip_amount", "3.791", "3.867", "0.00", "3.030", "222.000"],
                ["total_amount", "28.648", "20.752", "1.01", "21.900", "623.000"],
                ["passenger_count", "1.248", "0.640", "0.00", "1.000", "8.000"],
            ], [3.6, 2.4, 2.6, 2.2, 2.4, 2.3]),
        ],
        after=[
            "평균 속도와 상관이 큰 변수는 trip_distance(0.674), total_amount(0.535), "
            "fare_amount(0.532) 순이었습니다. 장거리 트립은 고속도로를 이용해 빠르고 단거리 "
            "트립은 도심 신호에 걸려 느리기 때문입니다. 시간대별로는 가장 느린 15시가 "
            "8.16 mph, 가장 빠른 5시가 19.21 mph 로 2.35배 차이가 났습니다.",
        ],
    ),
]


def chapter_results(doc):
    h(doc, "2. 실행 결과", 1, page_break=True)
    body(doc,
         "python main.py 를 실행하면 7단계가 순서대로 수행됩니다. 각 단계의 실제 실행 화면과 "
         "그때 산출된 수치를 함께 정리했습니다. 아래 화면은 모두 같은 한 번의 실행에서 나온 "
         "것으로, 총 소요 시간은 15.3초였습니다.")

    for st in STAGES:
        h(doc, f"{st['num']}. {st['title']}", 2, page_break=True)
        for t in st["paras"]:
            body(doc, t)
        for header, rows, widths in st["tables"]:
            make_table(doc, header, rows, widths=widths)
        for t in st.get("after", []):
            body(doc, t)
        sub(doc, "실행 결과")
        cap(doc, st["cap"], st["label"])

    # ── 2-4 시각화 ──
    h(doc, "2-4. 4단계. 시각화 — Seaborn 정적 / Plotly 인터랙티브", 2, page_break=True)
    body(doc,
         "모든 차트에 제목과 축 레이블을 포함했고, 한글이 깨지지 않도록 설치된 폰트를 확인해 "
         "명시적으로 적용했습니다. Plotly 차트는 화면 출력에 그치지 않고 write_html() 로 "
         "파일까지 저장했습니다.")
    make_table(doc, ["차트", "도구", "유형", "파일"], [
        ["시간대 × 요일별 평균 속도", "Seaborn", "히트맵 (그룹 비교)", "01_speed_heatmap_seaborn.png"],
        ["변수 간 상관계수", "Seaborn", "히트맵 (상관관계)", "02_correlation_heatmap_seaborn.png"],
        ["정체 여부별 주행거리 분포", "Seaborn", "히스토그램 (분포)", "03_distance_distribution_seaborn.png"],
        ["시간대별 속도와 정체 비율", "Plotly", "이중축 (인터랙티브)", "04_hourly_speed_plotly.html"],
        ["정체율 상위 출발 존", "Plotly", "막대 (인터랙티브)", "05_zone_jam_rate_plotly.html"],
    ], widths=[4.6, 1.8, 3.6, 5.5])
    cap(doc, "cap_04_visualize.png", "[ 4단계 실행 — 차트 5종 생성 및 저장 ]")

    sub(doc, "Seaborn 정적 차트")
    fig(doc, "01_speed_heatmap_seaborn.png",
        "[ 그림 1 ] 시간대·요일별 평균 주행 속도 — 평일 낮 시간대가 뚜렷한 적색 띠를 이룬다")
    fig(doc, "02_correlation_heatmap_seaborn.png",
        "[ 그림 2 ] 수치형 변수 간 피어슨 상관계수")
    fig(doc, "03_distance_distribution_seaborn.png",
        "[ 그림 3 ] 정체 여부에 따른 주행거리 분포 — 정체 트립이 단거리에 몰려 있다")

    sub(doc, "Plotly 인터랙티브 차트")
    body(doc,
         "아래는 저장된 HTML 을 브라우저에서 연 화면입니다. 실제 파일에서는 마우스를 올리면 "
         "해당 시간대의 트립 수·속도·정체 비율이 툴팁으로 표시되고, 범례 클릭으로 계열을 "
         "켜고 끌 수 있습니다.")
    cap(doc, "cap_04_hourly_speed_plotly.png",
        "[ 그림 4 ] 시간대별 평균 속도와 정체 비율 (이중축, 인터랙티브)")
    cap(doc, "cap_05_zone_jam_rate_plotly.png",
        "[ 그림 5 ] 정체율 상위 출발 존 20개 (색상 = 평균 속도, 인터랙티브)")

    # ── 2-5 통계 ──
    h(doc, "2-5. 5단계. 통계 분석 — t-test 및 p-value 해석", 2, page_break=True)
    body(doc,
         "표본 크기가 다르고 등분산을 가정할 수 없으므로 Welch's t-test(equal_var=False)를 "
         "사용했습니다. 유의수준은 α = 0.05 입니다.")
    make_table(doc, ["검정 대상", "집단 A 평균", "집단 B 평균", "평균 차이", "t", "p-value", "판정"], [
        ["정체 여부에 따른 주행거리(mi)", "1.388 (정체)", "3.927 (비정체)", "−2.539", "−721.303", "≈ 0", "유의"],
        ["주말/평일 평균속도(mph)", "10.736 (주말)", "9.420 (평일)", "+1.315", "170.996", "≈ 0", "유의"],
        ["출퇴근/그 외 평균속도(mph)", "9.250 (출퇴근)", "10.143 (그 외)", "−0.892", "−132.667", "≈ 0", "유의"],
    ], widths=[4.2, 2.3, 2.3, 1.8, 1.9, 1.4, 1.6], size=8)
    body(doc,
         "세 검정 모두 p < α 이므로 귀무가설을 기각했습니다. 즉 정체 트립은 유의하게 짧고, "
         "주말은 평일보다 유의하게 빠르며, 출퇴근 시간대는 유의하게 느립니다. 세 검정의 해석 "
         "규칙이 동일하므로 common.interpret_p() 하나로 공유해 문구를 자동 생성했습니다.")
    sub(doc, "실행 결과")
    cap(doc, "cap_05_ttest.png", "[ 5단계 실행 — 귀무가설·검정통계량·p-value 및 해석 문구 출력 ]")

    # ── 2-6 ML ──
    h(doc, "2-6. 6단계. ML Pipeline — 전처리 + 모델 학습 및 저장", 2, page_break=True)
    body(doc,
         "전처리와 모델을 하나의 Pipeline 객체로 묶었습니다. 전처리를 Pipeline 안에 두면 "
         "학습 데이터에서만 구한 통계가 예측 시 그대로 적용되어, 테스트 데이터 정보가 학습에 "
         "새어드는 것을 막을 수 있습니다.")
    add_code(doc, """Pipeline
 ├─ preprocess : ColumnTransformer
 │   ├─ numeric     (5개) : SimpleImputer(median)        → StandardScaler
 │   └─ categorical (4개) : SimpleImputer(most_frequent) → OneHotEncoder
 └─ classifier : RandomForestClassifier""")
    make_table(doc, None, [
        ["수치형 피처", "hour, dow, is_weekend, trip_distance, passenger_count"],
        ["범주형 피처", "PULocationID, DOLocationID, VendorID, RatecodeID"],
        ["제외 (누수)", "duration_min, speed_mph, fare_amount, tip_amount, total_amount"],
    ], widths=[3.0, 12.5], bold_first=True)
    body(doc,
         "존 ID 는 정수로 저장돼 있지만 숫자 크기에 의미가 없는 명목형이라 범주형으로 "
         "취급했습니다. 학습 표본은 200,000행을 층화 추출해 8:2 로 나눴고, 기준선으로 "
         "DummyClassifier(most_frequent)를 같은 Pipeline 구조로 함께 학습해 개선폭을 "
         "측정했습니다.")
    make_table(doc, ["타깃", "정확도", "F1(macro)", "기준선 F1", "F1 개선", "학습 시간", "모델 파일"], [
        ["jam_abs 절대정체", "0.7964", "0.7651", "0.3966", "+0.3685", "3.6초", "54.9 MB"],
        ["jam_rel 상대정체", "0.7339", "0.6575", "0.4002", "+0.2573", "3.8초", "56.4 MB"],
    ], widths=[3.4, 2.0, 2.2, 2.2, 2.0, 1.9, 1.8])
    body(doc,
         "학습된 파이프라인은 joblib.dump() 로 outputs/models/ 에 저장했습니다. 전처리기가 "
         "Pipeline 안에 포함돼 있어 저장 파일 하나만 불러오면 원본 형태의 입력을 그대로 "
         "예측에 쓸 수 있습니다.")
    sub(doc, "실행 결과")
    cap(doc, "cap_06_model.png",
        "[ 6-1 ] 두 타깃의 지표 비교 · 분류 리포트 · 혼동 행렬")
    cap(doc, "cap_06b_model_save.png",
        "[ 6-2 · 6-3 ] joblib 모델 저장과 두 타깃 정의 비교")

    # ── 2-7 자동화 ──
    h(doc, "2-7. 7단계. 자동화 — report.md 자동 생성", 2, page_break=True)
    body(doc,
         "각 단계가 반환한 결과를 context 딕셔너리에 누적하고, report.py 가 이를 마크다운으로 "
         "조립합니다. 외부 의존성(tabulate) 없이 표를 만들기 위해 common.to_markdown_table() "
         "을 직접 구현했습니다.")
    make_table(doc, ["산출물", "개수 / 크기"], [
        ["outputs/report.md", "378줄 / 16.6 KB"],
        ["정적 차트 (PNG)", "3개"],
        ["인터랙티브 차트 (HTML)", "2개"],
        ["학습된 모델 (pkl)", "2개"],
        ["분석 표 (CSV)", "5개"],
        ["총 실행 시간", "15.3초"],
    ], widths=[6.0, 9.5])
    sub(doc, "실행 결과")
    cap(doc, "cap_07_report.png", "[ 7단계 실행 — report.md 생성 및 전체 산출물 요약 ]")


OPINIONS = [
    ("결측 진단 — 개수 집계가 아닌 구조 판별", [
        "처음에는 isna().sum() 으로 컬럼별 결측 수를 세고 넘어가려 했습니다. 그런데 결측이 "
        "있는 5개 컬럼의 결측 수가 955,371건으로 전부 똑같은 것이 눈에 띄었습니다. 결측 "
        "위치를 불리언 마스크로 비교해 보니 .equals() 가 True, 즉 완전히 같은 행들이었습니다.",
        "더 확인해보니 이 블록 안팎의 VendorID 구성이 겹치지 않았습니다. 값이 개별적으로 빠진 "
        "것이 아니라 서로 다른 데이터 파이프라인이 한 파일에 섞여 있던 것입니다. 결측 처리 "
        "방법을 정하기 전에 결측이 무작위인지 구조적인지부터 판별해야 한다는 것을 배웠습니다.",
    ]),
    ("isna() 로 잡히지 않는 코드값 결측", [
        "NYC TLC 데이터는 ‘알 수 없음’을 NaN 이 아니라 약속된 코드값으로 기록합니다. "
        "RatecodeID = 99 가 14만 건(3.44%), 존 ID 264/265 가 2.5만 건(0.62%)이었는데 isna() "
        "는 이걸 전혀 잡지 못했습니다. Day 1 실습에서 Adult 데이터의 결측이 NaN 이 아니라 "
        "'?' 로 들어왔던 것과 같은 구조였습니다.",
        "처리 방침은 컬럼마다 다르게 정했습니다. 존 ID 는 모델의 핵심 피처라 미상이면 쓸 수 "
        "없어 제거했고, RatecodeID = 99 는 ‘미상’이라는 사실 자체가 정보이고 보조 피처라서 "
        "별도 범주로 남겼습니다. 덕분에 14만 행을 버리지 않아도 됐습니다.",
    ]),
    ("Pandas 와 Polars 의 집계 결과 차이 — null 그룹 처리", [
        "같은 집계를 두 방식으로 수행하자 그룹 수가 8개와 9개로 달랐습니다. Pandas 의 "
        "groupby 는 그룹 키가 결측인 행을 기본적으로 제외하지만, Polars 의 group_by 는 null 을 "
        "하나의 그룹으로 유지하기 때문이었고, 그 차이가 879,400건이었습니다.",
        "Pandas 만 썼다면 88만 건이 집계에서 빠진 것을 알아채지 못했을 것입니다. 두 도구를 "
        "함께 쓰는 실익이 속도가 아니라 이런 기본 동작 차이를 드러내는 데 있다는 것을 알게 "
        "됐습니다. 덧붙여 단순 로딩만 할 때는 Lazy 가 크게 빨랐는데 집계까지 붙이니 격차가 "
        "좁혀졌습니다. Lazy 의 이점은 컬럼을 적게 읽는 데서 오는 것이지 연산 자체가 빠른 것이 "
        "아니라는 점도 수치로 확인했습니다.",
    ]),
    ("높은 성능과 좋은 가설의 구분 — 타깃 누수", [
        "가설을 처음부터 하나로 정하지 않았습니다. 후보 타깃을 여러 개 만들어 실제로 끝까지 "
        "학습시켜 보고, 그 결과를 보고 최종 가설을 골랐습니다. 머리로 고르는 것보다 돌려 보는 "
        "편이 빨랐기 때문인데, 파이프라인을 먼저 만들어 둔 덕분에 타깃만 바꿔 끼우면 됐습니다.",
        "그 과정에서 “공항 트립인가”(Airport_fee > 0)를 예측했더니 F1 이 0.99 가 나왔습니다. "
        "처음에는 잘 나온 결과로 보았으나, 수치가 지나치게 높아 원인을 확인했습니다.",
        "Airport_fee 는 목적지가 아니라 출발지가 공항일 때 붙는 요금이었습니다. "
        "PULocationID == 132(JFK)면 95.2%, 138(LGA)면 99.3%, 그 외 존이면 0.27% 가 "
        "Airport_fee > 0 이었습니다. 모델이 예측한 게 아니라 존 번호 룩업 테이블을 외운 "
        "것이었고, 전형적인 타깃 누수라 이 가설은 폐기했습니다.",
        "이 경험 때문에 최종 모델에서는 누수 차단을 명시적으로 설계했습니다. duration_min 은 "
        "속도 계산의 분모라서, fare_amount 는 NYC 요금이 거리+시간 기반이라 소요시간이 새어 "
        "들어와서 제외했습니다. config.py 의 피처 정의에 “승차 시점에 알 수 있는 정보만”이라는 "
        "원칙을 주석으로 명시해 나중에도 실수하지 않도록 했습니다.",
    ]),
    ("타깃 정의에 따른 성능과 해석의 교환", [
        "절대정체(F1 0.7651)가 상대정체(0.6575)보다 점수가 높습니다. 그런데 각 타깃이 실제로 "
        "어떤 트립을 정체로 분류하는지 보니 성격이 달랐습니다. 절대정체로 분류된 트립의 "
        "39.2%가 1마일도 안 됐습니다.",
        "맨해튼 도심 크로스타운은 교통 상황과 무관하게 원래 느립니다. 즉 절대정체는 “교통 "
        "정체”가 아니라 상당 부분 “이게 도심 단거리 트립인지”를 맞히고 있었습니다. 피처를 "
        "하나씩 빼며 확인해보니 출발/도착 존만으로도 F1 0.64 가 나와서 이 추정이 맞았습니다.",
        "상대정체는 같은 경로의 평소 속도를 기준으로 삼아 이 성격을 통제합니다. 점수는 낮지만 "
        "“평소보다 막혔다”는 의미에 훨씬 충실합니다. 점수만 보고 절대정체를 골랐다면 겉보기엔 "
        "좋지만 실제로는 다른 것을 예측하는 모델을 만들었을 것입니다. 그래서 두 정의를 모두 "
        "학습해 보고서에서 비교하는 방식을 택했습니다.",
    ]),
    ("기본 하이퍼파라미터의 재검토 — 모델 크기와 성능", [
        "RandomForestClassifier 를 기본 설정으로 학습했더니 모델 파일이 636MB 가 나왔습니다. "
        "트리를 끝까지 쪼개기 때문입니다. min_samples_leaf=5 로 바꾸니 F1 이 0.7606 → 0.7651 "
        "로 오히려 개선되면서 크기는 55MB 로 줄었고, 학습 시간도 20초대에서 4초 안팎으로 "
        "짧아졌습니다.",
        "과적합을 줄이는 설정이 성능과 용량을 동시에 개선한 경우라, 기본값을 그대로 쓰기 전에 "
        "산출물의 크기까지 확인해야 한다는 것을 배웠습니다.",
    ]),
]


def chapter_opinion(doc):
    h(doc, "3. 코드 분석 결과 및 본인 의견", 1, page_break=True)
    body(doc,
         "구현과 측정을 진행하면서 예상과 달랐던 지점들이 있었습니다. 그때마다 원인을 "
         "확인하고 코드를 고친 과정을 정리했습니다.")
    for title, paras in OPINIONS:
        sub(doc, title)
        for t in paras:
            body(doc, t)


def chapter_team(doc):
    h(doc, "4. 팀 의견", 1, page_break=True)
    body(doc,
         "팀에서는 데이터셋을 하나로 통일하고, 가설은 각자 세워서 따로 돌린 뒤 결과를 모아 "
         "비교하기로 했습니다. 그 과정에서 팀에 맞춰야 했던 것들을 정리합니다.")

    sub(doc, "데이터셋 통일과 가설의 개별 설정")
    body(doc,
         "팀에서 정한 것은 데이터셋 하나뿐이었습니다. “무엇을 예측할 것인가”는 각자 정하기로 "
         "했습니다. 정해진 타깃이 없는 데이터라 무엇을 예측 대상으로 삼느냐에 따라 필요한 "
         "컬럼도, 정제 범위도, 결과 해석도 달라지기 때문입니다.")
    body(doc,
         "각자 가설을 세워 끝까지 돌려 본 뒤, 결과를 모아 함께 비교하는 시간을 가졌습니다. "
         "데이터셋이 같으니 전처리에서 마주치는 문제(구조적 결측 블록, 코드값 결측)는 "
         "공통이었지만, 그 다음부터는 타깃 정의에 따라 쓰는 컬럼과 남기는 행이 서로 "
         "달라졌습니다. 같은 데이터에서 출발해도 문제를 어떻게 정의하느냐에 따라 파이프라인이 "
         "달라진다는 것을 서로의 코드로 확인한 셈입니다. 비교를 마친 뒤에는 소스코드를 하나로 "
         "통일했습니다.")

    sub(doc, "데이터셋 선정 — 세 후보 비교")
    body(doc, "세 데이터셋을 모두 로딩해 규모·결측률·분류 타깃 유무를 비교했습니다.")
    make_table(doc, ["항목", "NYC Taxi", "Stack Overflow 2024", "Adult Census"], [
        ["규모", "4,090,836행 × 20열", "65,437행 × 114열", "32,561행 × 15열"],
        ["결측", "5개 열, 셀 5.8%", "109개 열, 셀 38.8%", "3개 열, 셀 0.9%"],
        ["분류 타깃", "없음 (파생 필요)", "없음 (파생 필요)", "내장 (income)"],
    ], widths=[2.6, 4.4, 4.5, 4.0])
    body(doc,
         "Adult 는 타깃이 내장돼 있어 가장 안전한 선택이었지만, 팀에서는 NYC Taxi 를 "
         "택했습니다. 규모가 커서 Pandas·Polars 비교의 의미가 살고, 파생 타깃을 직접 정의하는 "
         "과정 자체가 분석의 내용이 되기 때문입니다. Stack Overflow 는 114개 열 중 109개에 "
         "결측이 있어 정제에만 시간이 다 들어갈 것으로 판단해 제외했습니다.")

    sub(doc, "정제 범위와 베이스라인 설정 — 이견이 있었던 부분")
    body(doc,
         "“데이터를 통일한다”고 했을 때 정제까지 똑같이 해야 하는지를 두고 이야기가 나왔습니다. "
         "팀에서는 전처리를 먼저 해서, 명백한 오류로 보이는 것은 빼고 공통 기준을 맞추자는 "
         "쪽이었습니다. 저는 결측을 미리 빼는 것에 다른 생각이 있었습니다. 어떤 데이터가 "
         "필요한지를 사람이 미리 정해도 되는지, 미리 빼면 그 판단이 옳았는지 확인할 방법까지 "
         "함께 없어지는 것은 아닌지가 걸렸습니다.")
    body(doc,
         "어느 쪽이 옳다고 정하기보다 그래서 어떻게 처리할지를 두고 이야기를 나눴고, "
         "베이스라인은 같게 주어 출발점을 맞추되 그 제거가 자기 가설에서도 타당했는지는 각자 "
         "실험으로 확인해 보자는 쪽으로 정리했습니다. 정제 범위를 처음부터 못 박지 말고 돌려 본 "
         "결과로 되짚어 보자는 이야기였습니다.")
    make_table(doc, ["층", "범위", "이유"], [
        ["공통 층", "어떤 가설에서도 오류값인 것만 제거 (음수 요금, 거리 0, 소요시간 0, 기간 밖 레코드)",
         "결측이라기보다 데이터 오류에 가깝다고 본 것들"],
        ["개인 층", "각자 가설이 사용하는 컬럼의 결측만 처리",
         "가설마다 필요한 컬럼이 달라서"],
    ], widths=[2.2, 7.3, 6.0])
    body(doc,
         "실제로 어떤 컬럼을 쓰느냐에 따라 쓸 수 있는 행 수가 크게 달라집니다. 955,371행 "
         "결측 블록이 통째로 걸리기 때문입니다.")
    make_table(doc, ["가설이 쓰는 컬럼", "결측 블록", "사용 가능 행"], [
        ["passenger_count 를 쓰지 않는 경우", "제거 불필요", "4,090,836 (100%)"],
        ["passenger_count 를 쓰는 경우", "전량 제거", "3,135,465 (76.6%)"],
    ], widths=[6.5, 3.5, 5.5])
    body(doc,
         "본 프로젝트는 정체 예측에 passenger_count 를 피처로 포함했기 때문에 이 블록을 제거하고 "
         "진행했습니다(2장의 정제 이력 첫 줄). 같은 정체 예측이라도 이 컬럼을 빼면 95만 행을 "
         "그대로 쓸 수 있으니, 어떤 컬럼을 쓰느냐가 곧 정제 범위를 정하는 셈이었습니다. "
         "그래서 정제를 먼저 끝내고 가설을 정하기보다 가설을 정한 뒤에 결측을 다루는 편이 "
         "낫지 않겠냐는 이야기를 나눴고, 그 내용을 공유 문서에 적어 두었습니다.")
    sub(doc, "제거 결과 대조 — 공통 clean 과 본 파이프라인")
    body(doc,
         "팀에서 만든 공통 clean 데이터는 4,090,836행에서 3,021,962행이 되었고(26.1% 제거), "
         "결측은 하나도 남지 않았습니다. 무엇이 빠졌는지 원본과 대조해 보니 벤더 두 곳이 "
         "통째로 사라져 있었습니다. 그리고 제 파이프라인은 clean 파일을 쓰지 않고 원본을 직접 "
         "읽는데도, 결과는 같았습니다.")
    make_table(doc, ["VendorID", "원본", "공통 clean", "제 파이프라인", "빠진 지점"], [
        ["1", "805,221", "661,811", "552,882", "일부 행만 제거"],
        ["2", "3,226,222", "2,360,151", "2,288,671", "일부 행만 제거"],
        ["6", "7,643", "0", "0", "전량이 결측 블록 안에 있었음"],
        ["7", "51,750", "0", "0", "전량이 소요시간 1분 미만이었음"],
    ], widths=[2.0, 2.7, 2.7, 2.9, 5.2], size=8)
    body(doc,
         "벤더 6 은 결측 블록과 완전히 겹쳐 결측 제거 단계에서 빠졌습니다. 벤더 7 은 결측이 "
         "하나도 없는데도 사라졌는데, 51,750건 전부가 소요시간 1분 미만이라 이상치 규칙에 전량 "
         "걸렸기 때문입니다. 한 벤더의 기록이 예외 없이 1분 미만인 이유가 무엇인지까지는 "
         "확인하지 못했고, 이런 경우를 어떻게 다뤄야 할지를 두고 이야기를 나눴습니다.")
    body(doc,
         "팀의 기준이 잘못됐다는 이야기는 아닙니다. 저도 같은 규칙을 쓰고 있었기 때문에 같은 "
         "결과가 나왔습니다. 다만 clean 데이터만 놓고 보면 결측도 이상치도 없어서 이런 일이 "
         "있었는지 알기 어렵고 원본과 대조해야 드러난다는 점을, 확인해 보고 나서야 알게 "
         "됐습니다.")
    body(doc,
         "벤더 7 을 살릴지, 살린다면 1분 미만 기록을 어떻게 다룰지는 앞으로 더 이야기해 볼 "
         "부분으로 남겨 두었습니다. 어디까지를 ‘잘못된 값’으로 보고 어디부터를 ‘다르게 기록된 "
         "값’으로 볼지도 같이 볼 문제라, 후속 과목인 피처 엔지니어링에서 다시 다뤄 보려 "
         "합니다.")

    sub(doc, "평가 지표 기준의 정리 및 공유")
    body(doc,
         "가설이 다르면 타깃의 클래스 비율도 달라서, 정확도만으로는 서로의 결과를 나란히 놓고 "
         "보기 어려웠습니다. 클래스가 불균형하면 아무것도 학습하지 않은 모델도 정확도가 높게 "
         "나오기 때문입니다. 실제로 본 프로젝트의 기준선(DummyClassifier)은 아무것도 배우지 "
         "않고도 정확도 0.657 이 나왔습니다.")
    body(doc,
         "그래서 기준선 모델을 함께 학습시켜 두고, 기준선 대비 F1(macro) 개선폭을 같이 보도록 "
         "공유 문서에 적어 두었습니다. 절대적인 점수가 아니라 “아무것도 안 한 것보다 얼마나 "
         "나아졌는가”를 봐야 가설이 서로 달라도 비교가 되기 때문입니다.")
    make_table(doc, ["보는 방식", "jam_abs", "jam_rel", "판단"], [
        ["정확도만", "0.7964", "0.7339", "절대정체가 나아 보임"],
        ["기준선 F1", "0.3966", "0.4002", "출발점이 서로 다름"],
        ["F1(macro) 개선폭", "+0.3685", "+0.2573", "가설이 달라도 비교 가능"],
    ], widths=[4.2, 3.4, 3.4, 4.5])
    body(doc,
         "여기에 더해 “F1 이 0.9 를 넘으면 타깃 누수를 의심하라”는 점검 기준도 같이 적어 뒀습니다. "
         "실제로 후보 가설 중 하나가 F1 0.99 로 나왔다가 누수로 확인돼 폐기된 사례가 있었기 "
         "때문입니다.")

    sub(doc, "조원 공유 문서 작성")
    body(doc,
         "각자 다른 가설을 돌리다 보니 어떤 컬럼을 어떻게 썼는지 서로 알기 어려웠습니다. "
         "그래서 docs/컬럼_모델_정리.md 를 만들어 원본 20개 컬럼을 “모델 피처 / 통계 대상 / "
         "타깃 계산 / 미사용”으로 분류하고, 누수 때문에 일부러 뺀 컬럼과 그 이유를 "
         "명시했습니다. 특히 “F1 이 0.9를 넘으면 누수를 의심하라”는 체크 포인트를 "
         "공유했습니다.")


def chapter_improve(doc):
    h(doc, "5. 추가 의견 — 개선 사항", 1, page_break=True)
    make_table(doc, ["구분", "현재", "개선 방향"], [
        ["외부 데이터", "택시 운행 기록만 사용", "날씨·사고·공사 정보를 결합해 같은 경로 내 변동 설명"],
        ["정체 기준", "하위 33% 분위수로 고정", "절대 속도(5mph) 또는 경로별 하위 10% 등 기준 민감도 분석"],
        ["기간", "단일 월(2026-05)", "여러 달 결합으로 계절성 반영 + Polars 와일드카드 스캔 활용"],
        ["모델 튜닝", "기본 하이퍼파라미터", "GridSearchCV 를 Pipeline 에 결합 (구조상 몇 줄이면 가능)"],
        ["평가", "단일 홀드아웃 8:2", "교차검증으로 성능 추정 안정화"],
        ["모델 비교", "RandomForest 단일", "LightGBM·로지스틱 회귀 등과 대조해 모델 선택 근거 확보"],
        ["산출물 크기", "모델 pkl 합계 111MB", "트리 수 축소 또는 모델 압축으로 저장소 부담 완화"],
        ["자동화", "수동 실행", "GitHub Actions 로 ruff·실행 검증 자동화"],
    ], widths=[2.4, 4.6, 8.5])

    h(doc, "6. 코드 품질 측면에서 신경 쓴 점", 1, page_break=True)
    make_table(doc, ["항목", "내용"], [
        ["설정 단일화", "경로·상수·컬럼 목록을 config.py 한곳에 모아 관리"],
        ["누수 방지 설계", "피처 정의에 “승차 시점에 알 수 있는 정보만” 원칙을 주석으로 명시"],
        ["전처리를 Pipeline 안에", "학습 데이터에서만 통계를 구하도록 해 테스트 정보 누수 차단"],
        ["중복 로직 제거", "interpret_p()(p-value 해석), section()(구분선), _drop_rows()(제거+이력) 공용화"],
        ["예외 처리 구분", "전체 실행을 막을 오류는 중단, 표 저장 실패 등은 경고 후 계속 진행"],
        ["재현성", "random_state=42 고정, 산출물에 생성 일시 기록"],
        ["문서화", "전 파일 머리말 주석(프로그램명·작성자·설명·변경 이력·실행 방법·산출 파일) + 모든 함수 docstring"],
        ["정적 검사", "ruff check (line-length 100, select E/F/I/UP/B) 통과"],
        ["대용량 파일 관리", ".gitignore 로 원본 데이터·모델 제외, 보고서가 참조하는 차트·표는 추적"],
    ], widths=[3.6, 11.9])

    sub(doc, "정리 과정에서 발견한 잠재 버그")
    body(doc,
         "컬럼 목록이 config.py, preprocess.py, data_loader.py 세 곳에 흩어져 있어 한곳으로 "
         "모으는 작업을 했는데, 그 과정에서 tip_amount 가 기술통계 대상에는 있는데 결측 검사 "
         "목록에는 빠져 있는 것을 발견했습니다. 지금 데이터는 tip_amount 결측이 0건이라 증상이 "
         "나타나지 않았지만, 결측이 있는 달을 넣으면 NaN 이 그대로 기술통계와 상관계수에 "
         "흘러들어갔을 것입니다. 흩어진 설정은 그 자체로 잠재 버그라는 것을 확인한 "
         "사례였습니다.")

    sub(doc, "라이브러리 간 설정 충돌 — 차트 한글 깨짐")
    body(doc,
         "차트의 한글이 계속 깨져서 원인을 찾아보니, sns.set_theme() 이 plt.rcParams 의 폰트 "
         "설정을 덮어쓰고 있었습니다. 테마를 먼저 적용한 뒤 폰트를 지정하는 순서로 바꿔 "
         "해결했고, 설치된 폰트를 실제로 확인한 뒤 적용하도록 _apply_korean_font() 함수로 "
         "분리해 다른 환경에서도 동작하게 했습니다.")
    add_code(doc, '''# 주의: sns.set_theme()이 rcParams의 폰트 설정을 덮어쓰므로 반드시 테마를 먼저 적용한다.
sns.set_theme(style="whitegrid")

def _apply_korean_font() -> str:
    """설치된 한글 폰트 중 사용 가능한 것을 찾아 matplotlib에 적용한다."""
    installed = {f.name for f in font_manager.fontManager.ttflist}
    for candidate in KOREAN_FONT_CANDIDATES:
        if candidate in installed:
            plt.rcParams["font.family"] = candidate
            plt.rcParams["axes.unicode_minus"] = False   # 마이너스 기호 깨짐 방지
            return candidate
    return ""''')

    sub(doc, "마무리")
    body(doc,
         "python main.py 한 번으로 로딩부터 보고서 생성까지 약 15초에 끝나도록 만든 것이 이번 "
         "실습에서 가장 만족스러운 부분입니다. 반복 실행이 빨라야 타깃 정의를 바꿔보고, 피처를 "
         "빼보고, 기준값을 조정해보는 시도를 부담 없이 할 수 있었습니다.")


def build():
    doc = Document(TEMPLATE)

    # 템플릿의 본문은 비우고 서식(스타일·용지 설정)만 물려받는다
    b = doc.element.body
    for child in list(b):
        if child.tag != qn("w:sectPr"):
            b.remove(child)
    for rid, rel in list(doc.part.rels.items()):
        if "image" in rel.reltype:
            del doc.part.rels[rid]

    # ── 표지 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(130)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(TITLE)
    r.font.size = Pt(18)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(SUBTITLE)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(AUTHOR)
    r.font.size = Pt(13)
    r.font.bold = True

    # ── 목차 ──
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("목차")
    r.font.size = Pt(18)
    r.font.bold = True
    add_toc(doc)

    chapter_overview(doc)
    chapter_results(doc)
    chapter_opinion(doc)
    chapter_team(doc)
    chapter_improve(doc)

    doc.save(OUT)
    print(f"생성 완료: {OUT}")
    if MISSING:
        print("누락된 이미지:")
        for m in MISSING:
            print("  -", m)
    else:
        print("모든 이미지가 삽입되었습니다.")


if __name__ == "__main__":
    build()
